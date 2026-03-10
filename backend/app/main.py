import time
import io
import logging
from pathlib import Path

from docx import Document
from docxtpl import DocxTemplate

from fastapi import FastAPI, HTTPException, File, Form, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response

from . import template_filler, retrieval, generation
from .models import QueryRequest, QueryResponse, SourceDocument


app = FastAPI(title="CFIR Generator API")
logger = logging.getLogger(__name__)

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:5500",
        "http://127.0.0.1:5500",
        "http://localhost:8000",
        "http://127.0.0.1:8000",
        "null",  # file:// origin during local testing
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.post("/api/query", response_model=QueryResponse)
async def query(request: QueryRequest) -> QueryResponse:
    print("✅ Real query handler called")
    start_time = time.time()
    try:
        # 1. Retrieve relevant chunks from Chroma
        chunks_with_scores = retrieval.retrieve_relevant_chunks(request.prompt, k=5)

        if not chunks_with_scores:
            context = "No relevant documents found in the database."
            sources = []
        else:
            context_parts = []
            sources = []
            for doc, score in chunks_with_scores:
                full_path = doc.metadata.get("source", "Unknown")
                filename = Path(full_path).name
                context_parts.append(f"[Source: {filename}]\n{doc.page_content}")
                sources.append(SourceDocument(
                    filename=filename,
                    snippet=doc.page_content[:200] + "...",
                    relevance_score=float(score)
                ))
            context = "\n\n---\n\n".join(context_parts)

        # 2. Generate CFIR using Ollama
        answer = generation.generate_cfir(
            query=request.prompt,
            context=context,
            model=request.model,
            temperature=request.temperature
        )

        processing_time = int((time.time() - start_time) * 1000)
        return QueryResponse(
            answer=answer,
            sources=sources,
            processing_time_ms=processing_time,
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/health")
async def health() -> dict:
    return {"status": "healthy", "version": "1.0.0"}


@app.post("/api/ingest", status_code=202)
async def ingest() -> dict:
    """Trigger OneDrive download and rebuild vector store for future real queries."""
    from .ingestion import download_files_from_onedrive

    download_files_from_onedrive()
    retrieval.create_vector_store()
    return {"message": "Ingestion started. Check logs."}


@app.post("/api/fill-template")
async def fill_template(
    file: UploadFile = File(...),
    prompt: str = Form(...),
):
    try:
        # 1. Retrieve context
        chunks_with_scores = retrieval.retrieve_relevant_chunks(prompt, k=5)
        if not chunks_with_scores:
            context = "No relevant documents found."
        else:
            context = "\n\n---\n\n".join([doc.page_content for doc, _ in chunks_with_scores])

        # 2. Define fields (customize as needed)
        required_fields = [
            "incident_id", "date", "type", "description",
            "impact", "actions_taken", "reporter_name",
            "amount_lost", "currency",
        ]

        # 3. AI extracts structured data
        extracted_data = generation.extract_structured_data(prompt, context, required_fields)
        if not isinstance(extracted_data, dict):
            extracted_data = {}
        logger.info("Extracted data keys: %s", list(extracted_data.keys()))

        # 4. Handle based on file type
        template_bytes = await file.read()
        filename = file.filename or "template"
        lower_filename = filename.lower()

        if lower_filename.endswith('.docx'):
            # Try Jinja2 first (in case template has {{placeholders}})
            template_stream = io.BytesIO(template_bytes)
            doc = DocxTemplate(template_stream)
            doc.render(extracted_data)
            output_stream = io.BytesIO()
            doc.save(output_stream)
            output_stream.seek(0)
            rendered = output_stream.getvalue()

            # Fall back if output still has placeholders or original had no Jinja placeholders.
            original_doc = Document(io.BytesIO(template_bytes))
            original_has_jinja = any('{{' in para.text for para in original_doc.paragraphs)
            if not original_has_jinja:
                for table in original_doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if any('{{' in para.text for para in cell.paragraphs):
                                original_has_jinja = True
                                break
                        if original_has_jinja:
                            break
                    if original_has_jinja:
                        break

            temp_doc = Document(io.BytesIO(rendered))
            unfilled = any('{{' in para.text for para in temp_doc.paragraphs)
            if not unfilled:
                for table in temp_doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if any('{{' in para.text for para in cell.paragraphs):
                                unfilled = True
                                break
                        if unfilled:
                            break
                    if unfilled:
                        break

            if unfilled or not original_has_jinja:
                # For non-Jinja templates, use AI field mapping first, then heuristic fallback.
                if not original_has_jinja:
                    template_text_parts = [para.text for para in original_doc.paragraphs]
                    for table in original_doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                template_text_parts.extend([para.text for para in cell.paragraphs])
                    template_text = "\n".join([text for text in template_text_parts if text])
                    mapping = generation.map_template_fields(template_text, required_fields)
                    logger.info("Template mapping size: %s", len(mapping) if isinstance(mapping, dict) else 0)
                    mapped = template_filler.fill_docx_with_mapping(template_bytes, extracted_data, mapping)
                    filled = template_filler.fill_docx_heuristic(mapped, extracted_data)
                else:
                    # Preserve any Jinja substitutions, then fill remaining blanks/checkboxes.
                    filled = template_filler.fill_docx_heuristic(rendered, extracted_data)
            else:
                filled = rendered

            if filled == template_bytes:
                logger.warning("Filled DOCX bytes unchanged for %s", filename)

            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        elif lower_filename.endswith('.pdf'):
            try:
                filled = template_filler.fill_pdf_form(template_bytes, extracted_data)
                media_type = "application/pdf"
            except Exception as e:
                logger.exception("PDF fill failed for %s: %s", filename, e)
                raise HTTPException(
                    status_code=400,
                    detail=(
                        f"Could not fill PDF form: {str(e)}. "
                        "Please use a fillable PDF form or DOCX template."
                    ),
                )
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type. Please upload .docx or .pdf")

        return Response(
            content=filled,
            media_type=media_type,
            headers={"Content-Disposition": f"attachment; filename=filled_{filename}"},
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("Unexpected fill-template failure: %s", e)
        raise HTTPException(status_code=500, detail="Template filling failed due to an internal error")
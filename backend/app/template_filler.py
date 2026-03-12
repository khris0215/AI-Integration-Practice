import io
import logging
import re

import pdfrw
from docx import Document

logger = logging.getLogger(__name__)


LABEL_TO_FIELD = {
    "name": "reporter_name",
    "reporter name": "reporter_name",
    "full name": "reporter_name",
    "reported by": "reporter_name",
    "department": "department",
    "contact number": "contact_number",
    "phone": "contact_number",
    "mobile": "contact_number",
    "email": "email",
    "email address": "email",
    "date": "date",
    "date of incident": "date",
    "incident date": "date",
    "time": "time",
    "time of incident": "time",
    "location": "location",
    "system": "system",
    "system affected": "system",
    "location/system affected": "system",
    "actions taken": "actions_taken",
    "action taken": "actions_taken",
    "recommendations": "recommendations",
    "recommended next actions": "recommendations",
    "incident id": "incident_id",
    "incident number": "incident_id",
    "type": "type",
    "fraud type": "type",
    "description": "description",
    "incident description": "description",
    "full narrative": "description",
    "impact": "impact",
    "amount lost": "amount_lost",
    "loss amount": "amount_lost",
    "currency": "currency",
    "evidence": "evidence_list",
    "evidence list": "evidence_list",
    "reporter information": "reporter_name",
    "incident information": "incident_id",
}

BLANK_PATTERN = re.compile(r"_{3,}")
UNRESOLVED_PATTERN = re.compile(r"\{\{[^{}]+\}\}|_{3,}")
CHECKBOX_PATTERN = re.compile(r"^\s*[☐☑☒□]\s*(.+?)\s*$")
INLINE_CHECKBOX_PATTERN = re.compile(r"[☐☑☒□]\s*([^☐☑☒□\n]+)")
BRACKET_CHECKBOX_PATTERN = re.compile(r"(\[\s*[xX ]?\s*\]|\(\s*[xX ]?\s*\))\s*")
STRICT_NO_INFER_FIELDS = {
    "reporter_name",
    "department",
    "contact_number",
    "email",
    "time",
    "location",
    "system",
    "amount_lost",
    "currency",
}

FRAUD_OPTION_LABELS = [
    "Phishing",
    "Business Email Compromise",
    "Unauthorized Transfer",
    "Identity Theft",
    "Malware / Ransomware",
    "Other",
]


def _canonical_fraud_label(incident_type: str) -> str:
    text = (incident_type or "").lower()
    if "phishing" in text:
        return "Phishing"
    if "business email compromise" in text or "bec" in text or "deepfake" in text:
        return "Business Email Compromise"
    if "unauthorized transfer" in text or "wire transfer" in text or "fund transfer" in text:
        return "Unauthorized Transfer"
    if "identity theft" in text:
        return "Identity Theft"
    if "ransomware" in text or "malware" in text:
        return "Malware / Ransomware"
    return "Other"


def normalize_label(label: str) -> str:
    """Normalize text labels for dictionary and partial matching."""
    normalized = re.sub(r"[^a-zA-Z0-9\s/]", " ", (label or "").lower())
    normalized = re.sub(r"\s+", " ", normalized).strip()
    normalized = re.sub(r"^\d+[.)]\s*", "", normalized)
    return normalized


def map_label_to_field(label_text: str) -> str | None:
    """Map a label to a known extraction field using exact then partial matching."""
    normalized = normalize_label(label_text)
    if not normalized:
        return None

    direct = LABEL_TO_FIELD.get(normalized)
    if direct:
        return direct

    for key, field in LABEL_TO_FIELD.items():
        if key in normalized or normalized in key:
            return field

    return None


def _iter_all_paragraphs(doc: Document):
    def _iter_table_paragraphs(table):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
                for nested in cell.tables:
                    yield from _iter_table_paragraphs(nested)

    for paragraph in doc.paragraphs:
        yield paragraph

    for table in doc.tables:
        yield from _iter_table_paragraphs(table)

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            yield paragraph
        for table in section.header.tables:
            yield from _iter_table_paragraphs(table)

        for paragraph in section.footer.paragraphs:
            yield paragraph
        for table in section.footer.tables:
            yield from _iter_table_paragraphs(table)


def _best_heading_from_previous(previous_text: str) -> str:
    cleaned = (previous_text or "").strip()
    if not cleaned:
        return ""
    cleaned = re.sub(r"^\d+[.)]\s*", "", cleaned)
    return cleaned


def _looks_like_heading(text: str) -> bool:
    stripped = (text or "").strip()
    if not stripped:
        return False
    if stripped.endswith(":"):
        return True
    lowered = re.sub(r"[^a-zA-Z0-9\s]", " ", stripped.lower())
    lowered = re.sub(r"\s+", " ", lowered).strip()
    return any(key in lowered for key in [
        "reporter information",
        "incident information",
        "type of fraud",
        "incident description",
        "actions taken",
        "recommended next actions",
    ])


def extract_blanks_with_context(doc: Document) -> list:
    """
    Extract blank targets from all document paragraphs, including tables/headers/footers.
    Each entry contains paragraph reference, line index, and label/heading context.
    """
    blanks = []
    previous_non_blank = ""

    for paragraph in _iter_all_paragraphs(doc):
        text = paragraph.text or ""
        lines = text.split("\n") if text else [""]

        for line_index, line in enumerate(lines):
            if line.strip() and not BLANK_PATTERN.search(line):
                previous_non_blank = line.strip()
                if _looks_like_heading(line.strip()):
                    previous_non_blank = line.strip()

            matches = list(BLANK_PATTERN.finditer(line))
            if not matches:
                continue

            before_first_blank = line[: matches[0].start()].strip()
            heading = _best_heading_from_previous(previous_non_blank)

            # If line is mostly underscores, rely on previous heading.
            mostly_blank = bool(re.fullmatch(r"\s*_{3,}\s*", line))
            label = before_first_blank if before_first_blank and not mostly_blank else heading

            for blank_idx, _ in enumerate(matches):
                blanks.append({
                    "paragraph": paragraph,
                    "paragraph_text": text,
                    "line_index": line_index,
                    "blank_index": blank_idx,
                    "full_text": line,
                    "label": label,
                    "preceding_heading": heading,
                })

    return blanks


def _replace_nth_blank(line: str, blank_index: int, replacement: str) -> str:
    """Replace the Nth underscore-blank occurrence in a line."""
    current = -1

    def repl(match: re.Match) -> str:
        nonlocal current
        current += 1
        if current == blank_index:
            return replacement
        return match.group(0)

    return BLANK_PATTERN.sub(repl, line)


def fill_blanks_with_ai(doc: Document, context: str, extracted_data: dict) -> bool:
    """
    Fill blanks using dictionary mapping first, then LLM fallback for unresolved labels.
    Returns True if any blank was filled.
    """
    from . import generation

    filled_any = False
    blanks = extract_blanks_with_context(doc)
    logger.info("Detected %s blank target(s) in template", len(blanks))
    inference_cache = {}

    # Replace from right to left per paragraph-line so blank indexes remain stable.
    blanks.sort(key=lambda t: (id(t["paragraph"]), t["line_index"], -t["blank_index"]))

    for target in blanks:
        paragraph = target["paragraph"]
        line_index = target["line_index"]
        blank_index = target["blank_index"]

        current_text = paragraph.text or ""
        lines = current_text.split("\n") if current_text else [""]
        if line_index >= len(lines):
            continue

        line = lines[line_index]
        label = target.get("label") or ""
        heading = target.get("preceding_heading") or ""

        field_name = map_label_to_field(label)
        value = ""

        if field_name:
            candidate = extracted_data.get(field_name)
            if isinstance(candidate, str) and candidate.strip():
                value = candidate.strip()
            elif candidate is not None:
                value = str(candidate).strip()

            # Do not guess sensitive identity/contact/financial fields.
            if not value and field_name in STRICT_NO_INFER_FIELDS:
                value = "Not provided"

        if not value:
            ai_label = label or heading or "field value"
            if heading and label and heading.lower() not in label.lower():
                ai_label = f"{heading} - {label}"
            if ai_label in inference_cache:
                ai_value = inference_cache[ai_label]
            else:
                ai_value = generation.infer_value_from_context(ai_label, context)
                inference_cache[ai_label] = ai_value
            value = ai_value.strip()

        if not value:
            value = "Not provided"

        lines[line_index] = _replace_nth_blank(line, blank_index, value)
        paragraph.text = "\n".join(lines)
        filled_any = True

    return filled_any


def _normalize_tokens(value: str) -> set:
    cleaned = re.sub(r"[^a-zA-Z0-9\s]", " ", (value or "").lower())
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    if not cleaned:
        return set()
    return {token for token in cleaned.split(" ") if token}


def _select_checkbox_lines(doc: Document, context: str, extracted_data: dict) -> int:
    """Mark the best matching fraud type checkbox line as checked."""
    from . import generation

    incident_type = str((extracted_data or {}).get("type") or "").strip()
    if not incident_type:
        incident_type = generation.infer_value_from_context("Type of Fraud", context).strip()
    if not incident_type:
        return 0

    type_tokens = _normalize_tokens(incident_type)
    if not type_tokens:
        return 0

    changed = 0
    in_fraud_section = False
    section_options = []

    def parse_checkbox_lines(paragraph_text: str) -> list[tuple[int, int, str]]:
        options = []
        lines = (paragraph_text or "").split("\n")
        for idx, line in enumerate(lines):
            inline_matches = list(INLINE_CHECKBOX_PATTERN.finditer(line))
            if inline_matches:
                for occ_idx, match in enumerate(inline_matches):
                    options.append((idx, occ_idx, match.group(1).strip()))
                continue

            match = CHECKBOX_PATTERN.match(line)
            if match:
                options.append((idx, 0, match.group(1).strip()))
        return options

    def _replace_checkbox_symbol(line: str, target_occurrence: int, marker: str) -> str:
        current_occurrence = -1

        def repl(match: re.Match) -> str:
            nonlocal current_occurrence
            current_occurrence += 1
            option_text = match.group(1).strip()
            symbol = marker if current_occurrence == target_occurrence else match.group(0)[0]
            return f"{symbol} {option_text}"

        return INLINE_CHECKBOX_PATTERN.sub(repl, line)

    def update_checkbox_lines(paragraph, line_entries: list[tuple[int, int]], checked_pos: int) -> int:
        local_changes = 0
        lines = (paragraph.text or "").split("\n")
        for pos, (line_idx, occ_idx) in enumerate(line_entries):
            if line_idx >= len(lines):
                continue
            current = lines[line_idx]
            marker = "☑" if pos == checked_pos else "☐"
            updated = _replace_checkbox_symbol(current, occ_idx, marker)
            if updated != current:
                lines[line_idx] = updated
                local_changes += 1
        if local_changes:
            paragraph.text = "\n".join(lines)
        return local_changes

    def flush_section_options(options):
        local_changed = 0
        if not options:
            return local_changed

        best_idx = None
        best_score = -1
        for idx, (_, option_text, _, _) in enumerate(options):
            option_tokens = _normalize_tokens(option_text)
            overlap = option_tokens.intersection(type_tokens)
            score = len(overlap)

            # Strong hint if the full option phrase appears in inferred type.
            if option_text and option_text.lower() in incident_type.lower():
                score += 5

            if score > best_score:
                best_score = score
                best_idx = idx

        if best_idx is None or best_score <= 0:
            # Fallback to Other when type cannot be matched reliably.
            for idx, (_, option_text, _, _) in enumerate(options):
                if "other" in option_text.lower():
                    best_idx = idx
                    break

        grouped = {}
        for idx, (paragraph, option_text, line_idx, occ_idx) in enumerate(options):
            grouped.setdefault(id(paragraph), {"paragraph": paragraph, "entries": []})
            grouped[id(paragraph)]["entries"].append((idx, option_text, line_idx, occ_idx))

        for group in grouped.values():
            entries = group["entries"]
            paragraph = group["paragraph"]
            line_entries = [(line_idx, occ_idx) for _, _, line_idx, occ_idx in entries]
            checked_pos = 0
            for pos, (global_idx, _, _, _) in enumerate(entries):
                if global_idx == best_idx:
                    checked_pos = pos
                    break
            local_changed += update_checkbox_lines(paragraph, line_entries, checked_pos)

        return local_changed

    for paragraph in _iter_all_paragraphs(doc):
        text = paragraph.text or ""
        normalized = re.sub(r"\s+", " ", text.lower()).strip()

        if "type of fraud" in normalized:
            in_fraud_section = True
            continue

        if in_fraud_section and _looks_like_heading(text) and "type of fraud" not in normalized:
            changed += flush_section_options(section_options)
            section_options = []
            in_fraud_section = False

        if not in_fraud_section:
            continue

        line_matches = parse_checkbox_lines(text)
        if not line_matches:
            continue

        for line_idx, occ_idx, option_text in line_matches:
            section_options.append((paragraph, option_text, line_idx, occ_idx))

    if section_options:
        changed += flush_section_options(section_options)

    return changed


def _force_mark_fraud_option(doc: Document, incident_type: str) -> int:
    """Fallback: mark standard fraud options even when template checkbox style is unusual."""
    selected = _canonical_fraud_label(incident_type)
    changed = 0

    for paragraph in _iter_all_paragraphs(doc):
        text = paragraph.text or ""
        if not text.strip():
            continue

        updated = text
        touched_any_label = False

        for label in FRAUD_OPTION_LABELS:
            escaped = re.escape(label)
            marker = "☑" if label == selected else "☐"

            # Existing unicode checkbox markers.
            p_unicode = re.compile(rf"([☐☑☒□])\s*({escaped})", re.IGNORECASE)
            if p_unicode.search(updated):
                updated = p_unicode.sub(lambda m: f"{marker} {m.group(2)}", updated)
                touched_any_label = True

            # Existing [ ] / [X] / ( ) / (X) checkbox markers.
            p_bracket = re.compile(rf"(\[\s*[xX ]?\s*\]|\(\s*[xX ]?\s*\))\s*({escaped})", re.IGNORECASE)
            if p_bracket.search(updated):
                updated = p_bracket.sub(lambda m: f"{marker} {m.group(2)}", updated)
                touched_any_label = True

        # If fraud labels exist without markers, prepend markers deterministically.
        if not touched_any_label:
            lines = updated.split("\n")
            line_changed = False
            for idx, line in enumerate(lines):
                line_strip = line.strip()
                for label in FRAUD_OPTION_LABELS:
                    if re.fullmatch(rf"(?i){re.escape(label)}:?", line_strip):
                        marker = "☑" if label == selected else "☐"
                        suffix = ":" if line_strip.endswith(":") else ""
                        lines[idx] = f"{marker} {label}{suffix}"
                        line_changed = True
                        break
            if line_changed:
                updated = "\n".join(lines)

        if updated != text:
            paragraph.text = updated
            changed += 1

    return changed


def add_warning_paragraph(doc: Document, warning: str):
    """Append a warning paragraph to the end of the document."""
    doc.add_paragraph(warning)


def has_unfilled_placeholders(doc: Document) -> bool:
    for paragraph in _iter_all_paragraphs(doc):
        text = paragraph.text or ""
        if UNRESOLVED_PATTERN.search(text):
            return True
    return False


def fill_docx_intelligently(original_template_bytes: bytes, context: str, extracted_data: dict) -> bytes:
    """
    Main DOCX fill pipeline:
    1) Fill blanks by dictionary/AI inference.
    2) Add warning if unresolved blanks remain.
    3) Return valid DOCX bytes.
    """
    doc = Document(io.BytesIO(original_template_bytes))
    fill_blanks_with_ai(doc, context, extracted_data or {})
    checkbox_updates = _select_checkbox_lines(doc, context, extracted_data or {})
    if checkbox_updates == 0:
        fallback_updates = _force_mark_fraud_option(doc, str((extracted_data or {}).get("type") or ""))
        checkbox_updates += fallback_updates
    logger.info("Checkbox updates applied: %s", checkbox_updates)

    if has_unfilled_placeholders(doc):
        logger.warning("Template still has unresolved placeholders after intelligent fill")
        add_warning_paragraph(doc, "Warning: Some fields could not be filled automatically.")

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


def validate_docx(docx_bytes: bytes) -> bool:
    """Validate generated DOCX by reloading with python-docx."""
    try:
        Document(io.BytesIO(docx_bytes))
        return True
    except Exception as exc:
        logger.exception("DOCX validation failed: %s", exc)
        return False


def fill_pdf_form(template_bytes: bytes, data: dict) -> bytes:
    """Optional helper for fillable PDF forms."""
    try:
        pdf = pdfrw.PdfReader(fdata=template_bytes)
        for page in pdf.pages:
            annotations = page.get('/Annots')
            if not annotations:
                continue
            for annotation in annotations:
                field = annotation.get('/T')
                if not field:
                    continue
                field_name = field[1:-1]
                for key, value in (data or {}).items():
                    if key.lower() == field_name.lower():
                        annotation.update(pdfrw.PdfDict(V=pdfrw.PdfString(str(value))))
        output = io.BytesIO()
        pdfrw.PdfWriter().write(output, pdf)
        output.seek(0)
        return output.getvalue()
    except Exception as exc:
        logger.exception("Failed PDF form fill: %s", exc)
        raise
